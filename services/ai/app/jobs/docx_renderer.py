"""DOCX renderer for intelligence briefs.

Generates .docx attachments for email briefs using python-docx.
Clean formatting: Arial, tables, section headers.
"""
import logging
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

BRIEFS_DIR = Path(__file__).parent.parent.parent / "data" / "briefs"


def _ensure_dir():
    BRIEFS_DIR.mkdir(parents=True, exist_ok=True)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
    return h


def _add_para(doc, text, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _add_table_row(table, cells, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(text)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                if bold or header:
                    run.bold = True


def render_daily_docx(data: dict, path: str | Path | None = None) -> Path:
    """Render daily brief as .docx file.

    Args:
        data: Daily brief data from assemble_daily_data().
        path: Output path (default: data/briefs/daily_YYYY-MM-DD.docx).

    Returns:
        Path to the generated .docx file.
    """
    _ensure_dir()
    brief_date = data.get("date", date.today().isoformat())
    if path is None:
        path = BRIEFS_DIR / f"daily_{brief_date}.docx"
    path = Path(path)

    doc = Document()

    # Title
    title = doc.add_heading("CFTC Daily Brief", level=0)
    for run in title.runs:
        run.font.name = "Arial"
    _add_para(doc, data.get("date_display", brief_date), size=12, color=(100, 116, 139))

    # Section 1: What Changed
    _add_heading(doc, "What Changed Overnight", level=1)
    changes = data.get("what_changed", [])
    if changes:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for i, label in enumerate(["Type", "Change", "Time"]):
            table.rows[0].cells[i].text = label
            for p in table.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(9)

        for c in changes[:20]:
            _add_table_row(table, [
                c.get("entity_type", ""),
                c.get("summary", ""),
                (c.get("timestamp") or "")[:16],
            ])
    else:
        _add_para(doc, "No changes since last brief.", color=(148, 163, 184))

    # Section 2: Action List
    _add_heading(doc, "Action List", level=1)
    actions = data.get("action_list", [])
    if actions:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, label in enumerate(["Priority", "Title", "Matter", "Detail"]):
            table.rows[0].cells[i].text = label
            for p in table.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(9)

        for a in actions[:15]:
            _add_table_row(table, [
                a.get("tag", ""),
                a.get("title", ""),
                a.get("matter", ""),
                a.get("detail", ""),
            ])
    else:
        _add_para(doc, "No action items today.", color=(148, 163, 184))

    # Section 3: Meetings
    _add_heading(doc, "Today's Meetings", level=1)
    meetings = data.get("meetings", [])
    if meetings:
        for m in meetings:
            _add_para(doc, m.get("title", ""), bold=True, size=11)
            details = []
            if m.get("start_time"):
                details.append(m["start_time"][:5])
            if m.get("meeting_type"):
                details.append(m["meeting_type"])
            if m.get("location"):
                details.append(m["location"])
            if details:
                _add_para(doc, " \u2022 ".join(details), color=(100, 116, 139))

            participants = [p.get("full_name", p.get("name", "")) for p in m.get("participants", [])]
            if participants:
                _add_para(doc, "Participants: " + ", ".join(participants), size=9, color=(100, 116, 139))

            if m.get("prep_narrative"):
                p = doc.add_paragraph()
                p.style = "Quote" if "Quote" in [s.name for s in doc.styles] else None
                run = p.add_run(m["prep_narrative"])
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.italic = True
    else:
        _add_para(doc, "No meetings today.", color=(148, 163, 184))

    # Section 4: Follow-Ups
    _add_heading(doc, "Follow-Ups Due", level=1)
    followups = data.get("followups", [])
    if followups:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, label in enumerate(["Person", "Organization", "Due", "Purpose"]):
            table.rows[0].cells[i].text = label
            for p in table.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(9)

        for f in followups[:10]:
            _add_table_row(table, [
                f.get("name", ""),
                f.get("organization", ""),
                f.get("next_date", ""),
                f.get("purpose", "") or f.get("interaction_type", ""),
            ])
    else:
        _add_para(doc, "No follow-ups due.", color=(148, 163, 184))

    # Section 5: Team Pulse
    _add_heading(doc, "Team Pulse", level=1)
    pulse = data.get("team_pulse", {})
    overdue = pulse.get("overdue_count", 0)
    if overdue:
        by_assignee = pulse.get("overdue_by_assignee", {})
        _add_para(doc, f"{overdue} overdue tasks: " + ", ".join(f"{k} ({v})" for k, v in by_assignee.items()))
    overloaded = pulse.get("overloaded_people", [])
    if overloaded:
        names = ", ".join(f"{p['name']} ({p['task_count']})" for p in overloaded)
        _add_para(doc, f"Overloaded: {names}")
    if not overdue and not overloaded:
        _add_para(doc, "No team execution risks.", color=(34, 197, 94))

    # Save
    doc.save(str(path))
    logger.info("DOCX saved: %s", path)
    return path



def render_weekly_docx(data, path=None):
    """Render weekly brief as .docx file."""
    _ensure_dir()
    brief_date = data.get("date", date.today().isoformat())
    if path is None:
        path = BRIEFS_DIR / f"weekly_{brief_date}.docx"
    path = Path(path)

    doc = Document()

    # Title
    title = doc.add_heading("CFTC Weekly Brief", level=0)
    for run in title.runs:
        run.font.name = "Arial"
    _add_para(doc, data.get("date_display", brief_date), size=12, color=(100, 116, 139))

    # Calibration
    cal = data.get("calibration", {})
    _add_heading(doc, "What I Got Wrong", level=1)
    if cal.get("has_data"):
        score = cal.get("score", 0)
        _add_para(doc, f"Signal Quality: {score}%", bold=True, size=14)
        _add_para(doc, f"Materialized: {cal.get('materialized', 0)}, Resolved: {cal.get('resolved', 0)}, Still open: {cal.get('still_open', 0)}, Wrong: {cal.get('wrong', 0)}")
    else:
        _add_para(doc, cal.get("message", "No calibration data."), color=(148, 163, 184))

    # Executive Summary
    _add_heading(doc, "Executive Summary", level=1)
    summary = data.get("executive_summary")
    if summary:
        p = doc.add_paragraph()
        run = p.add_run(summary)
        run.font.name = "Arial"
        run.font.size = Pt(10)
    else:
        _add_para(doc, "Not generated.", color=(148, 163, 184))

    # Portfolio Health
    _add_heading(doc, "Portfolio Health", level=1)
    portfolio = data.get("portfolio", {})
    for posture, label in [("critical", "Critical"), ("important", "Important"), ("strategic", "Strategic"), ("monitoring", "Monitoring")]:
        items = portfolio.get(posture, [])
        if items:
            _add_para(doc, f"{label} ({len(items)})", bold=True, size=11)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for i, h in enumerate(["Matter", "Status", "Deadline", "Owner"]):
                table.rows[0].cells[i].text = h
                for p in table.rows[0].cells[i].paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.name = "Arial"
                        run.font.size = Pt(9)
            for m in items:
                _add_table_row(table, [m.get("title", ""), m.get("status", ""), m.get("nearest_deadline", ""), m.get("next_step_owner", "")])

    # Decision Docket
    _add_heading(doc, "Decision Docket", level=1)
    decisions = data.get("decisions", [])
    if decisions:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["Decision", "Matter", "Owner", "Due"]):
            table.rows[0].cells[i].text = h
            for p in table.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
        for d in decisions:
            _add_table_row(table, [d.get("title", ""), d.get("matter_title", ""), d.get("decision_owner", ""), d.get("due_date", "")])
    else:
        _add_para(doc, "No open decisions.", color=(148, 163, 184))

    # Team View
    _add_heading(doc, "Team Management", level=1)
    team = data.get("team", {})
    workload = team.get("workload", [])
    if workload:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["Person", "Tasks", "Matters", "Overdue"]):
            table.rows[0].cells[i].text = h
            for p in table.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
        for w in workload:
            _add_table_row(table, [w.get("name", ""), str(w.get("open_tasks", 0)), str(w.get("open_matters", 0)), str(w.get("overdue", 0))])
    drifting = team.get("drifting_matters", [])
    if drifting:
        _add_para(doc, "Drifting Matters:", bold=True)
        for dm in drifting:
            _add_para(doc, f"  {dm.get('title', '')} - {dm.get('days_stale', 0)} days stale - {dm.get('owner', '')}", size=9)

    # Stakeholders
    _add_heading(doc, "Stakeholders & Relationships", level=1)
    stak = data.get("stakeholders", {})
    touchpoints = stak.get("touchpoints_due", [])
    if touchpoints:
        _add_para(doc, "Touchpoints Due:", bold=True)
        for tp in touchpoints:
            _add_para(doc, f"  {tp.get('name', '')} ({tp.get('organization', '')}) - {tp.get('next_date', '')} - {tp.get('purpose', '')}", size=9)
    neglected = stak.get("neglected", [])
    if neglected:
        _add_para(doc, "Neglected Relationships:", bold=True)
        for n in neglected:
            _add_para(doc, f"  {n.get('name', '')} ({n.get('category', '')}) - {n.get('days_since', 0)} days", size=9)
    if not touchpoints and not neglected:
        _add_para(doc, "No stakeholder actions needed.", color=(148, 163, 184))

    # Deadlines
    _add_heading(doc, "Deadlines & Horizon Scan", level=1)
    dl_data = data.get("deadlines", {})
    for period, label in [("two_weeks", "Next 2 Weeks"), ("thirty_days", "Next 30 Days"), ("ninety_days", "Next 90 Days")]:
        items = dl_data.get(period, [])
        if items:
            _add_para(doc, f"{label} ({len(items)}):", bold=True)
            for item in items:
                _add_para(doc, f"  {item.get('date', '')} - {item.get('deadline_type', '')} - {item.get('matter_title', '')} - {item.get('owner', '')}", size=9)

    # Data Hygiene
    _add_heading(doc, "Data Hygiene", level=1)
    hygiene = data.get("hygiene", {})
    _add_para(doc, f"Tracker Health Score: {hygiene.get('score', 0)}%", bold=True, size=14)
    for c in hygiene.get("checks", []):
        field = c.get("field", "")
        _add_para(doc, f"  {field}: {c['count']}/{c['total']} ({c['pct']}%)", size=9)

    doc.save(str(path))
    logger.info("Weekly DOCX saved: %s", path)
    return path
