"""Briefing Document Export Service.

Generates professional Word documents (.docx) from AI-analyzed comment data.
Uses python-docx for document generation.
"""

import json
import os
import logging
from datetime import datetime
from typing import List, Optional

import anthropic
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy import select, func, desc
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.models import Comment, CommentTag, TagType

logger = logging.getLogger(__name__)


def _add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2E75B6')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')  # 10pt
    rPr.append(sz)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _setup_styles(doc):
    """Configure document styles for a professional CFTC briefing."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Century Schoolbook'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    title_style = doc.styles['Title']
    title_style.font.name = 'Century Schoolbook'
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    title_style.paragraph_format.space_after = Pt(4)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Century Schoolbook'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Century Schoolbook'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2d, 0x5a, 0xa0)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Century Schoolbook'
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    return doc


def _add_table_row(table, cells_data, bold=False, bg_color=None):
    """Add a row to a table with optional formatting."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(text) if text else '—'
        for paragraph in cell.paragraphs:
            paragraph.style.font.size = Pt(9)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Century Schoolbook'
                if bold:
                    run.bold = True
        if bg_color:
            from docx.oxml.ns import qn
            shading = cell._element.get_or_add_tcPr()
            shading_elem = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): bg_color,
                qn('w:val'): 'clear',
            })
            shading.append(shading_elem)
    return row


async def _generate_statutory_analysis(tier1_comments: list) -> Optional[dict]:
    """Use Claude Opus to synthesize statutory interpretation disputes across Tier 1 comments."""
    if not settings.ANTHROPIC_API_KEY or settings.ANTHROPIC_API_KEY == "your_anthropic_key_here":
        return None

    # Build a condensed summary of all Tier 1 legal arguments
    summaries = []
    for c in tier1_comments:
        structured = c.ai_summary_structured or {}
        name = c.commenter_name or c.commenter_organization or 'Anonymous'
        legal = structured.get('legal_challenges', [])
        key_args = structured.get('key_arguments', [])
        position = c.sentiment or 'Unknown'

        if not legal and not key_args:
            continue

        entry = f"COMMENTER: {name} | POSITION: {position}\n"
        if key_args:
            entry += "KEY ARGUMENTS:\n"
            for arg in key_args:
                entry += f"  - {arg.get('topic', '')}: {'; '.join(arg.get('sub_points', []))}\n"
                if arg.get('requested_action'):
                    entry += f"    Requested: {arg['requested_action']}\n"
        if legal:
            entry += "LEGAL CITATIONS:\n"
            for lc in legal:
                entry += f"  - {lc.get('citation', '')}: {lc.get('theory', '')}\n"
        summaries.append(entry)

    if not summaries:
        return None

    combined = "\n---\n".join(summaries)
    # Truncate if too long
    if len(combined) > 60000:
        combined = combined[:60000] + "\n...[truncated]"

    system_prompt = """You are a senior legal analyst at the CFTC preparing a briefing for the Deputy General Counsel.

Analyze the Tier 1 comment summaries below and produce a JSON object identifying the key STATUTORY INTERPRETATION DISPUTES in this rulemaking. Focus on:

1. Disputed statutory terms — what specific words or phrases in the Commodity Exchange Act are being contested
2. Competing interpretations — who argues what each term means, and why
3. The legal basis for each interpretation (textualism, legislative history, prior CFTC interpretation, case law)
4. How the post-Chevron/Loper Bright landscape affects each dispute

Return ONLY a JSON object with this structure:
{
    "overview": "2-3 paragraph executive summary of the major statutory interpretation battleground in this rulemaking",
    "disputes": [
        {
            "statutory_provision": "The specific statute section (e.g., CEA § 5c(c)(5)(C))",
            "disputed_term": "The specific word or phrase in dispute (e.g., 'gaming', 'involve')",
            "commission_position": "How the proposed rule interprets the term",
            "challenger_position": "How opponents interpret the term",
            "key_commenters_challenging": ["Commenter A", "Commenter B"],
            "key_commenters_supporting": ["Commenter C"],
            "legal_basis_for_challenge": "The legal reasoning challengers use",
            "legal_basis_for_commission": "The legal reasoning supporting the Commission's reading",
            "risk_assessment": "How strong is the challenge — HIGH/MEDIUM/LOW risk to the final rule"
        }
    ],
    "loper_bright_implications": "How the end of Chevron deference specifically affects this rulemaking's legal exposure"
}"""

    try:
        client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-opus-4-20250514",
            max_tokens=4000,
            system=system_prompt,
            messages=[{"role": "user", "content": f"Analyze statutory interpretation disputes across these Tier 1 comments:\n\n{combined}"}],
        )

        text = response.content[0].text.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1]
            if text.endswith("```"):
                text = text[:-3]
            text = text.strip()
        return json.loads(text)

    except Exception as e:
        logger.error(f"Error generating statutory analysis: {e}")
        return None


async def _generate_legal_authority_analysis(tier1_comments: list, top_citations: list) -> Optional[list]:
    """Use Claude to explain each cited legal authority and its relevance."""
    if not settings.ANTHROPIC_API_KEY or settings.ANTHROPIC_API_KEY == "your_anthropic_key_here":
        return None
    if not top_citations:
        return None

    # Gather context from Tier 1 comments about how these citations are used
    cite_context = []
    for c in tier1_comments:
        structured = c.ai_summary_structured or {}
        legal = structured.get('legal_challenges', [])
        name = c.commenter_name or c.commenter_organization or 'Anonymous'
        for lc in legal:
            cite_context.append(f"{name} ({c.sentiment}): {lc.get('citation', '')} — {lc.get('theory', '')}")

    citations_list = "\n".join([f"- {cite} (cited {count} times)" for cite, count in top_citations])
    context_text = "\n".join(cite_context[:100])

    system_prompt = """You are a senior legal analyst at the CFTC preparing a briefing for the Deputy General Counsel.

For each cited legal authority, provide a JSON array where each element has:
{
    "citation": "The case or statute name exactly as provided",
    "times_cited": <number>,
    "type": "case" | "statute" | "regulation" | "executive_order" | "other",
    "summary": "1-2 sentence explanation of what this authority holds or establishes",
    "relevance": "2-3 sentences explaining why commenters are citing this authority in this rulemaking and what argument it supports",
    "cited_by": "Which side cites it — 'challengers', 'supporters', or 'both'",
    "risk_note": "Brief note on how this authority affects legal risk to the final rule"
}

Return ONLY a JSON array. Be precise and legally accurate."""

    try:
        client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            system=system_prompt,
            messages=[{"role": "user", "content": f"Explain these cited legal authorities:\n\n{citations_list}\n\nContext from comments:\n{context_text}"}],
        )

        text = response.content[0].text.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1]
            if text.endswith("```"):
                text = text[:-3]
            text = text.strip()
        return json.loads(text)

    except Exception as e:
        logger.error(f"Error generating legal authority analysis: {e}")
        return None


async def _generate_comment_narrative(tier1_comments: list, tier2_comments: list) -> Optional[str]:
    """Synthesize a narrative summary of what the major commenters said."""
    if not settings.ANTHROPIC_API_KEY or settings.ANTHROPIC_API_KEY == "your_anthropic_key_here":
        return None

    summaries = []
    for c in tier1_comments + tier2_comments[:20]:
        structured = c.ai_summary_structured or {}
        name = c.commenter_name or c.commenter_organization or 'Anonymous'
        exec_sum = structured.get('executive_summary') or c.ai_summary or ''
        position = c.sentiment or 'Unknown'
        if exec_sum:
            summaries.append(f"COMMENTER: {name} | POSITION: {position}\n{exec_sum[:500]}")

    if not summaries:
        return None

    combined = "\n---\n".join(summaries)
    if len(combined) > 50000:
        combined = combined[:50000] + "\n...[truncated]"

    system_prompt = """You are a senior legal analyst at the CFTC writing a narrative synthesis of public comments for the Deputy General Counsel.

Write a comprehensive narrative (4-6 paragraphs) that tells the story of this comment record. Do NOT use bullet points or lists. Write in flowing prose suitable for a legal memorandum. The narrative should:

1. Open with the overall landscape — how many major commenters weighed in, what the general split was
2. Describe the main arguments of the opposition — name specific commenters and what they argued
3. Describe the main arguments of supporters — name specific commenters and what they argued
4. Identify areas of consensus or unexpected agreement across positions
5. Close with an assessment of the most significant legal and policy challenges the Commission must address in the final rule

Name specific commenters throughout. Write as though briefing a senior attorney who needs to understand the full picture quickly."""

    try:
        client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-opus-4-20250514",
            max_tokens=3000,
            system=system_prompt,
            messages=[{"role": "user", "content": f"Synthesize a narrative of these comments:\n\n{combined}"}],
        )
        return response.content[0].text.strip()
    except Exception as e:
        logger.error(f"Error generating comment narrative: {e}")
        return None


async def _generate_draft_preamble(tier1_comments: list, stat_analysis: dict, docket_number: str) -> Optional[str]:
    """Generate a draft preamble 'Response to Comments' section."""
    if not settings.ANTHROPIC_API_KEY or settings.ANTHROPIC_API_KEY == "your_anthropic_key_here":
        return None

    # Build condensed comment summaries grouped by topic
    comment_data = []
    for c in tier1_comments:
        structured = c.ai_summary_structured or {}
        name = c.commenter_name or c.commenter_organization or 'Anonymous'
        position = c.sentiment or 'Unknown'
        key_args = structured.get('key_arguments', [])
        requested_changes = structured.get('requested_changes', [])

        entry = f"COMMENTER: {name} ({position})\n"
        for arg in key_args:
            entry += f"  ARGUMENT: {arg.get('topic', '')}\n"
            for sp in arg.get('sub_points', []):
                entry += f"    - {sp}\n"
            if arg.get('requested_action'):
                entry += f"  REQUESTED: {arg['requested_action']}\n"
        for rc in requested_changes:
            entry += f"  REQUESTED CHANGE: {rc}\n"
        comment_data.append(entry)

    if not comment_data:
        return None

    combined = "\n---\n".join(comment_data)
    if len(combined) > 50000:
        combined = combined[:50000]

    # Include statutory disputes if available
    disputes_text = ""
    if stat_analysis and stat_analysis.get('disputes'):
        for d in stat_analysis['disputes']:
            disputes_text += f"\nDISPUTED TERM: \"{d.get('disputed_term', '')}\"\n"
            disputes_text += f"  Commission position: {d.get('commission_position', '')}\n"
            disputes_text += f"  Challenger position: {d.get('challenger_position', '')}\n"

    system_prompt = """You are a senior attorney at the CFTC Office of General Counsel drafting the "Response to Comments" section of a final rule preamble for the Federal Register.

Write in the formal style of a Federal Register preamble. Use the standard structure:

1. Group comments by topic/issue
2. For each issue:
   - Summarize the comments received (naming specific commenters)
   - State "The Commission has considered these comments and..."
   - Provide the Commission's response, explaining the reasoning
   - Note any changes made to the final rule in response, OR explain why the Commission is not making changes

Use phrases like:
- "Several commenters, including [X] and [Y], argued that..."
- "The Commission disagrees with commenters who suggested..."
- "In response to comments from [X], the Commission has modified..."
- "The Commission appreciates these comments but notes that..."
- "After careful consideration, the Commission has determined..."

Write 8-12 paragraphs of substantive response. This is a DRAFT for attorney review — be thorough but note where the Commission's position needs to be confirmed with brackets like [Commission to confirm position on X].

Do NOT use bullet points. Write in flowing Federal Register prose."""

    try:
        client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-opus-4-20250514",
            max_tokens=6000,
            system=system_prompt,
            messages=[{"role": "user", "content": f"Draft a preamble response to these comments for docket {docket_number}:\n\n{combined}\n\nKey statutory disputes:\n{disputes_text}"}],
        )
        return response.content[0].text.strip()
    except Exception as e:
        logger.error(f"Error generating draft preamble: {e}")
        return None


async def generate_briefing_doc(
    db: AsyncSession,
    docket_number: str,
    output_path: str,
) -> str:
    """Generate a comprehensive briefing document for a docket.
    
    Returns the path to the generated .docx file.
    """
    doc = Document()
    _setup_styles(doc)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ------------------------------------------------------------------
    # Fetch all data
    # ------------------------------------------------------------------

    # Stats
    stats_query = select(
        func.count(Comment.id).label('total'),
        func.count(Comment.id).filter(Comment.tier == 1).label('tier1'),
        func.count(Comment.id).filter(Comment.tier == 2).label('tier2'),
        func.count(Comment.id).filter(Comment.tier == 3).label('tier3'),
        func.count(Comment.id).filter(Comment.sentiment == 'SUPPORT').label('support'),
        func.count(Comment.id).filter(Comment.sentiment == 'OPPOSE').label('oppose'),
        func.count(Comment.id).filter(Comment.sentiment == 'MIXED').label('mixed'),
        func.count(Comment.id).filter(Comment.sentiment == 'NEUTRAL').label('neutral'),
        func.count(Comment.id).filter(Comment.is_form_letter == True).label('form_letters'),
    ).where(Comment.docket_number == docket_number)
    stats_result = await db.execute(stats_query)
    stats = stats_result.one()

    # Tier 1 comments
    t1_query = select(Comment).where(
        Comment.docket_number == docket_number,
        Comment.tier == 1,
    ).order_by(Comment.sentiment.asc(), desc(func.length(Comment.comment_text or '')))
    t1_result = await db.execute(t1_query)
    tier1_comments = t1_result.scalars().all()

    # Tier 2 comments
    t2_query = select(Comment).where(
        Comment.docket_number == docket_number,
        Comment.tier == 2,
    ).order_by(Comment.sentiment.asc())
    t2_result = await db.execute(t2_query)
    tier2_comments = t2_result.scalars().all()

    # Legal citations
    citations_query = select(
        CommentTag.tag_value,
        func.count(CommentTag.id).label('cite_count'),
    ).join(Comment).where(
        Comment.docket_number == docket_number,
        CommentTag.tag_type == TagType.LEGAL_CITATION,
    ).group_by(CommentTag.tag_value).order_by(desc('cite_count')).limit(25)
    citations_result = await db.execute(citations_query)
    top_citations = citations_result.all()

    # Topics
    topics_query = select(
        CommentTag.tag_value,
        func.count(CommentTag.id).label('topic_count'),
    ).join(Comment).where(
        Comment.docket_number == docket_number,
        CommentTag.tag_type == TagType.TOPIC,
    ).group_by(CommentTag.tag_value).order_by(desc('topic_count')).limit(20)
    topics_result = await db.execute(topics_query)
    top_topics = topics_result.all()

    # ------------------------------------------------------------------
    # Title Page
    # ------------------------------------------------------------------

    doc.add_paragraph('')  # spacer
    doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('COMMODITY FUTURES TRADING COMMISSION')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    run.font.name = 'Century Schoolbook'
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('COMMENT LETTER ANALYSIS')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    run.font.name = 'Century Schoolbook'
    run.bold = True

    docket_para = doc.add_paragraph()
    docket_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = docket_para.add_run(f'Docket: {docket_number}')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Century Schoolbook'

    doc.add_paragraph('')

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'Prepared: {datetime.now().strftime("%B %d, %Y")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Century Schoolbook'

    classification = doc.add_paragraph()
    classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = classification.add_run('PRIVILEGED — ATTORNEY WORK PRODUCT')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)
    run.font.name = 'Century Schoolbook'
    run.bold = True

    doc.add_page_break()

    # ------------------------------------------------------------------
    # Table of Contents placeholder
    # ------------------------------------------------------------------

    doc.add_heading('TABLE OF CONTENTS', level=1)
    toc_items = [
        'I. Executive Overview',
        'II. Sentiment Breakdown',
        'III. Legal Challenges Tracker',
        'IV. Statutory Interpretation Disputes',
        'V. Synthesis of Major Comments',
        'VI. Tier 1 Comment Summaries — Oppose',
        'VII. Tier 1 Comment Summaries — Support',
        'VIII. Tier 1 Comment Summaries — Mixed/Neutral',
        'IX. Tier 2 Summary Table',
        'X. Draft Preamble — Response to Comments',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ------------------------------------------------------------------
    # I. Executive Overview
    # ------------------------------------------------------------------

    doc.add_heading('I. EXECUTIVE OVERVIEW', level=1)

    total = stats.total
    overview_text = (
        f'The Commission received {total} public comment letters on {docket_number}. '
        f'Of these, {stats.tier1} were classified as Tier 1 (requiring deep analysis), '
        f'{stats.tier2} as Tier 2 (substantive), and {stats.tier3} as Tier 3 (standard). '
        f'{stats.form_letters} comments were identified as part of {len(set(c.form_letter_group_id for c in tier1_comments if c.form_letter_group_id))} '
        f'form letter campaigns.'
    )
    doc.add_paragraph(overview_text)

    sentiment_text = (
        f'Sentiment analysis reveals significant opposition to the proposed rule: '
        f'{stats.oppose} comments ({(stats.oppose/total*100):.1f}%) oppose, '
        f'{stats.support} ({(stats.support/total*100):.1f}%) support, '
        f'{stats.mixed} ({(stats.mixed/total*100):.1f}%) express mixed views, and '
        f'{stats.neutral} ({(stats.neutral/total*100):.1f}%) are neutral.'
    )
    doc.add_paragraph(sentiment_text)

    # Stats table
    stats_table = doc.add_table(rows=1, cols=4)
    stats_table.style = 'Light Grid Accent 1'
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = stats_table.rows[0].cells
    for i, label in enumerate(['Category', 'Count', '% of Total', 'Notes']):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    table_data = [
        ('Total Comments', total, '100%', ''),
        ('Tier 1 (Critical)', stats.tier1, f'{stats.tier1/total*100:.1f}%', 'Deep analysis required'),
        ('Tier 2 (Substantive)', stats.tier2, f'{stats.tier2/total*100:.1f}%', 'Executive summary'),
        ('Tier 3 (Standard)', stats.tier3, f'{stats.tier3/total*100:.1f}%', 'Brief summary'),
        ('Form Letters', stats.form_letters, f'{stats.form_letters/total*100:.1f}%', 'Identified campaigns'),
        ('Oppose', stats.oppose, f'{stats.oppose/total*100:.1f}%', ''),
        ('Support', stats.support, f'{stats.support/total*100:.1f}%', ''),
        ('Mixed', stats.mixed, f'{stats.mixed/total*100:.1f}%', ''),
        ('Neutral', stats.neutral, f'{stats.neutral/total*100:.1f}%', ''),
    ]
    for row_data in table_data:
        _add_table_row(stats_table, row_data)

    doc.add_paragraph('')

    # ------------------------------------------------------------------
    # II. Sentiment Breakdown
    # ------------------------------------------------------------------

    doc.add_heading('II. SENTIMENT BREAKDOWN', level=1)

    # Visual bar
    bar_text = (
        f'OPPOSE {"█" * int(stats.oppose/total*50)} {stats.oppose} ({stats.oppose/total*100:.1f}%)\n'
        f'SUPPORT {"█" * int(stats.support/total*50)} {stats.support} ({stats.support/total*100:.1f}%)\n'
        f'MIXED {"█" * int(stats.mixed/total*50)} {stats.mixed} ({stats.mixed/total*100:.1f}%)\n'
        f'NEUTRAL {"█" * int(stats.neutral/total*50)} {stats.neutral} ({stats.neutral/total*100:.1f}%)'
    )
    bar_para = doc.add_paragraph()
    run = bar_para.add_run(bar_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('')

    # ------------------------------------------------------------------
    # III. Legal Challenges Tracker
    # ------------------------------------------------------------------

    doc.add_heading('III. LEGAL CHALLENGES TRACKER', level=1)

    if top_citations:
        doc.add_paragraph(
            'The following legal authorities were most frequently cited by commenters. '
            'AI analysis provides context for each citation and its relevance to this rulemaking.'
        )

        # Generate AI explanations
        logger.info("Generating AI legal authority analysis...")
        legal_explanations = await _generate_legal_authority_analysis(tier1_comments, top_citations)

        if legal_explanations:
            for i, auth in enumerate(legal_explanations, 1):
                # Citation header
                doc.add_heading(
                    f'{i}. {auth.get("citation", "Unknown")} — cited {auth.get("times_cited", "?")} times',
                    level=2,
                )

                # Type badge
                auth_type = auth.get('type', 'other').upper()
                cited_by = auth.get('cited_by', 'unknown')
                type_para = doc.add_paragraph()
                type_run = type_para.add_run(f'Type: {auth_type}')
                type_run.font.size = Pt(9)
                type_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                type_run2 = type_para.add_run(f'  |  Cited by: {cited_by}')
                type_run2.font.size = Pt(9)
                type_run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                # Summary
                if auth.get('summary'):
                    p = doc.add_paragraph()
                    r = p.add_run('What it establishes: ')
                    r.bold = True
                    r.font.size = Pt(10)
                    r2 = p.add_run(auth['summary'])
                    r2.font.size = Pt(10)

                # Relevance
                if auth.get('relevance'):
                    p = doc.add_paragraph()
                    r = p.add_run('Why commenters cite it: ')
                    r.bold = True
                    r.font.size = Pt(10)
                    r2 = p.add_run(auth['relevance'])
                    r2.font.size = Pt(10)

                # Risk note
                if auth.get('risk_note'):
                    p = doc.add_paragraph()
                    r = p.add_run('Risk to final rule: ')
                    r.bold = True
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)
                    r2 = p.add_run(auth['risk_note'])
                    r2.font.size = Pt(10)

                doc.add_paragraph('')
        else:
            # Fallback to simple table if AI fails
            cite_table = doc.add_table(rows=1, cols=2)
            cite_table.style = 'Light Grid Accent 1'
            cite_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = cite_table.rows[0].cells
            hdr[0].text = 'Case / Statute'
            hdr[1].text = 'Times Cited'
            for p in hdr[0].paragraphs + hdr[1].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
            for citation, count in top_citations:
                _add_table_row(cite_table, [citation, count])
    else:
        doc.add_paragraph('No legal citations were extracted from the comment record.')

    doc.add_paragraph('')

    # ------------------------------------------------------------------
    # IV. Statutory Interpretation Disputes (AI-synthesized)
    # ------------------------------------------------------------------

    doc.add_heading('IV. STATUTORY INTERPRETATION DISPUTES', level=1)

    logger.info("Generating AI statutory interpretation analysis...")
    stat_analysis = await _generate_statutory_analysis(tier1_comments)

    if stat_analysis:
        # Overview
        overview = stat_analysis.get('overview', '')
        if overview:
            for para_text in overview.split('\n\n'):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    for run in p.runs:
                        run.font.size = Pt(11)

        # Individual disputes
        disputes = stat_analysis.get('disputes', [])
        for i, dispute in enumerate(disputes, 1):
            doc.add_heading(
                f'{i}. "{dispute.get("disputed_term", "Unknown")}" — {dispute.get("statutory_provision", "")}',
                level=2,
            )

            # Risk badge
            risk = dispute.get('risk_assessment', 'UNKNOWN')
            risk_para = doc.add_paragraph()
            risk_run = risk_para.add_run(f'Risk to Final Rule: {risk}')
            risk_run.bold = True
            risk_run.font.size = Pt(10)
            if 'HIGH' in risk.upper():
                risk_run.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)
            elif 'MEDIUM' in risk.upper():
                risk_run.font.color.rgb = RGBColor(0xcc, 0x88, 0x00)
            else:
                risk_run.font.color.rgb = RGBColor(0x22, 0x88, 0x22)

            # Commission's position
            if dispute.get('commission_position'):
                p = doc.add_paragraph()
                r = p.add_run("Commission's Interpretation: ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(dispute['commission_position'])
                r2.font.size = Pt(10)

            # Challenger's position
            if dispute.get('challenger_position'):
                p = doc.add_paragraph()
                r = p.add_run("Challenger's Interpretation: ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(dispute['challenger_position'])
                r2.font.size = Pt(10)

            # Legal basis for challenge
            if dispute.get('legal_basis_for_challenge'):
                p = doc.add_paragraph()
                r = p.add_run("Legal Basis for Challenge: ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(dispute['legal_basis_for_challenge'])
                r2.font.size = Pt(10)

            # Legal basis for Commission
            if dispute.get('legal_basis_for_commission'):
                p = doc.add_paragraph()
                r = p.add_run("Legal Basis for Commission: ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(dispute['legal_basis_for_commission'])
                r2.font.size = Pt(10)

            # Key commenters
            challengers = dispute.get('key_commenters_challenging', [])
            supporters = dispute.get('key_commenters_supporting', [])
            if challengers:
                p = doc.add_paragraph()
                r = p.add_run("Key Commenters Challenging: ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(", ".join(challengers))
                r2.font.size = Pt(10)
            if supporters:
                p = doc.add_paragraph()
                r = p.add_run("Key Commenters Supporting: ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(", ".join(supporters))
                r2.font.size = Pt(10)

            doc.add_paragraph('')

        # Loper Bright implications
        loper = stat_analysis.get('loper_bright_implications', '')
        if loper:
            doc.add_heading('Post-Chevron / Loper Bright Implications', level=2)
            p = doc.add_paragraph(loper)
            for run in p.runs:
                run.font.size = Pt(11)
    else:
        doc.add_paragraph(
            'Statutory interpretation analysis could not be generated. '
            'Review Tier 1 comment summaries below for individual interpretive arguments.'
        )

    doc.add_page_break()

    # ------------------------------------------------------------------
    # V. Synthesis of Major Comments
    # ------------------------------------------------------------------

    doc.add_heading('V. SYNTHESIS OF MAJOR COMMENTS', level=1)

    logger.info("Generating comment narrative synthesis...")
    narrative = await _generate_comment_narrative(tier1_comments, tier2_comments)

    if narrative:
        for para_text in narrative.split('\n\n'):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                for run in p.runs:
                    run.font.size = Pt(11)
    else:
        doc.add_paragraph('Comment narrative synthesis could not be generated.')

    doc.add_page_break()

    # ------------------------------------------------------------------
    # V-VII. Tier 1 Comment Summaries by Sentiment
    # ------------------------------------------------------------------

    sentiment_groups = [
        ('VI', 'OPPOSE', 'Oppose'),
        ('VII', 'SUPPORT', 'Support'),
        ('VIII', 'MIXED/NEUTRAL', 'Mixed/Neutral'),
    ]

    for section_num, sentiment_key, sentiment_label in sentiment_groups:
        doc.add_heading(f'{section_num}. TIER 1 COMMENT SUMMARIES — {sentiment_key}', level=1)

        if sentiment_key == 'MIXED/NEUTRAL':
            group = [c for c in tier1_comments if c.sentiment in ('MIXED', 'NEUTRAL', None)]
        else:
            group = [c for c in tier1_comments if c.sentiment == sentiment_key]

        if not group:
            doc.add_paragraph(f'No Tier 1 comments with {sentiment_label.lower()} sentiment.')
            continue

        for idx, comment in enumerate(group, 1):
            structured = comment.ai_summary_structured or {}
            name = comment.commenter_name or comment.commenter_organization or 'Anonymous'

            # Comment header
            doc.add_heading(f'{idx}. {name}', level=2)

            # Metadata line
            meta = doc.add_paragraph()
            meta_text = f'Document ID: {comment.document_id}'
            if comment.submission_date:
                meta_text += f'  |  Date: {comment.submission_date}'
            if comment.page_count:
                meta_text += f'  |  Pages: {comment.page_count}'
            meta_text += f'  |  Position: {comment.sentiment or "Unknown"}'
            if structured.get('commenter_type'):
                meta_text += f'  |  Type: {structured["commenter_type"]}'
            run = meta.add_run(meta_text)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            # Link to original comment
            link_para = doc.add_paragraph()
            # Build proper CFTC comment URL
            comment_id_match = (comment.document_id or '').replace('CFTC-COMMENT-', '')
            if comment_id_match:
                cftc_url = f'https://comments.cftc.gov/PublicComments/ViewComment.aspx?id={comment_id_match}'
                _add_hyperlink(link_para, 'View Comment on CFTC.gov', cftc_url)
                link_run = link_para.add_run('  |  ')
                link_run.font.size = Pt(9)
                link_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            if comment.regulations_gov_url and comment.regulations_gov_url.startswith('http'):
                _add_hyperlink(link_para, 'Download Original PDF', comment.regulations_gov_url)

            # Executive Summary
            exec_summary = structured.get('executive_summary') or comment.ai_summary
            if exec_summary:
                doc.add_heading('Executive Summary', level=3)
                # Prefix with commenter name for context
                p = doc.add_paragraph()
                name_run = p.add_run(f'{name} ')
                name_run.bold = True
                name_run.font.size = Pt(10)
                summary_run = p.add_run(exec_summary)
                summary_run.font.size = Pt(10)

            # Key Arguments
            key_args = structured.get('key_arguments', [])
            if key_args:
                doc.add_heading('Key Arguments', level=3)
                for i, arg in enumerate(key_args, 1):
                    # Argument topic
                    p = doc.add_paragraph()
                    run = p.add_run(f'{i}. {arg.get("topic", "Unnamed Argument")}')
                    run.bold = True
                    run.font.size = Pt(10)

                    # Sub-points
                    for sp in arg.get('sub_points', []):
                        bp = doc.add_paragraph(style='List Bullet')
                        bp_run = bp.add_run(sp)
                        bp_run.font.size = Pt(10)

                    # Requested action
                    if arg.get('requested_action'):
                        ra = doc.add_paragraph()
                        ra_label = ra.add_run('Requested Action: ')
                        ra_label.bold = True
                        ra_label.font.size = Pt(10)
                        ra_label.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
                        ra_text = ra.add_run(arg['requested_action'])
                        ra_text.font.size = Pt(10)

            # Legal Challenges
            legal = structured.get('legal_challenges', [])
            if legal:
                doc.add_heading('Legal Challenges Raised', level=3)
                for lc in legal:
                    p = doc.add_paragraph()
                    cite_run = p.add_run(f'{lc.get("citation", "Unknown")}: ')
                    cite_run.bold = True
                    cite_run.font.size = Pt(10)
                    theory_run = p.add_run(lc.get('theory', ''))
                    theory_run.font.size = Pt(10)

            # Data / Evidence
            data_ev = structured.get('data_evidence', [])
            if data_ev:
                doc.add_heading('Data & Evidence Provided', level=3)
                for d in data_ev:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp_run = bp.add_run(d)
                    bp_run.font.size = Pt(10)

            # Requested Changes
            changes = structured.get('requested_changes', [])
            if changes:
                doc.add_heading('Requested Specific Changes', level=3)
                for i, ch in enumerate(changes, 1):
                    p = doc.add_paragraph()
                    run = p.add_run(f'{i}. {ch}')
                    run.font.size = Pt(10)

            # Key Quotes
            quotes = structured.get('key_quotes', [])
            if quotes:
                doc.add_heading('Key Quotes', level=3)
                for q in quotes:
                    p = doc.add_paragraph()
                    run = p.add_run(f'"{q}"')
                    run.italic = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Add spacing between comments
            doc.add_paragraph('')

        doc.add_page_break()

    # ------------------------------------------------------------------
    # VIII. Tier 2 Summary Table
    # ------------------------------------------------------------------

    doc.add_heading('IX. TIER 2 SUMMARY TABLE', level=1)

    doc.add_paragraph(
        f'{len(tier2_comments)} Tier 2 comments received substantive executive summaries.'
    )

    if tier2_comments:
        t2_table = doc.add_table(rows=1, cols=4)
        t2_table.style = 'Light Grid Accent 1'
        t2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t2_table.rows[0].cells
        for i, label in enumerate(['Commenter', 'Position', 'Pages', 'Summary']):
            hdr[i].text = label
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        for c in tier2_comments:
            name = c.commenter_name or c.commenter_organization or 'Anonymous'
            summary = c.ai_summary or '—'
            if len(summary) > 250:
                summary = summary[:250] + '...'
            _add_table_row(t2_table, [
                name,
                c.sentiment or '—',
                c.page_count or '—',
                summary,
            ])

    # ------------------------------------------------------------------
    # X. Draft Preamble — Response to Comments
    # ------------------------------------------------------------------

    doc.add_page_break()
    doc.add_heading('X. DRAFT PREAMBLE — RESPONSE TO COMMENTS', level=1)

    disclaimer = doc.add_paragraph()
    d_run = disclaimer.add_run(
        'DRAFT FOR ATTORNEY REVIEW — This section was generated by AI and requires '
        'substantive review and revision before inclusion in any Federal Register document. '
        'Bracketed items [like this] indicate areas requiring Commission input or confirmation.'
    )
    d_run.font.size = Pt(9)
    d_run.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)
    d_run.italic = True

    doc.add_paragraph('')

    logger.info("Generating draft preamble...")
    # Use the statutory analysis we already generated (if available)
    draft_preamble = await _generate_draft_preamble(tier1_comments, stat_analysis, docket_number)

    if draft_preamble:
        for para_text in draft_preamble.split('\n\n'):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                for run in p.runs:
                    run.font.size = Pt(11)
    else:
        doc.add_paragraph('Draft preamble could not be generated.')

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------

    doc.save(output_path)
    logger.info(f"Briefing document saved to {output_path}")
    return output_path
